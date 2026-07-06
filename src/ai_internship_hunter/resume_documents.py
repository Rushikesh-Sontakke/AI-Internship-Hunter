from __future__ import annotations

import os
import re
from datetime import date
from html import escape
from pathlib import Path

from .resume import Experience, Project, ResumeSource, TailoredResume


# compact_reference_guide with named resume overrides:
# Letter portrait; 0.55-inch margins for a one-page application resume;
# Calibri/Helvetica 9.2-point body; restrained navy section hierarchy;
# proposal_centerpiece title block adapted to name, target headline, and contacts.
NAVY = "17365D"
BLUE = "1F4D78"
GRAY = "555555"

# Fallback maps <b>/<i> markup to the built-in Helvetica family.
_HELVETICA_FONTS = {
    "regular": "Helvetica", "bold": "Helvetica-Bold",
    "italic": "Helvetica-Oblique", "bolditalic": "Helvetica-BoldOblique",
}
_pdf_fonts_cache: dict[str, str] | None = None


def _slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.casefold()).strip("-")


def resume_file_stem(name: str) -> str:
    """Professional resume filename stem, e.g. ``Rushikesh_Sontakke_Resume``."""

    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")
    return f"{cleaned}_Resume" if cleaned else "Resume"


def cover_letter_file_stem(name: str) -> str:
    """Cover-letter filename stem, e.g. ``Rushikesh_Sontakke_Cover_Letter``."""

    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")
    return f"{cleaned}_Cover_Letter" if cleaned else "Cover_Letter"


def build_cover_letter_pdf(source: ResumeSource, body_text: str, path: Path) -> Path:
    """Render a one-page business-letter cover letter matching the resume style.

    ``body_text`` is the drafted letter (greeting through signature). It is laid
    out verbatim under a letterhead; no content is added or invented.
    """

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    fonts = _pdf_fonts()
    document = SimpleDocTemplate(
        str(path), pagesize=letter,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.8 * inch, bottomMargin=0.8 * inch,
        title=f"Cover Letter - {source.name}", author=source.name,
    )
    styles = getSampleStyleSheet()
    gray = colors.HexColor(f"#{GRAY}")
    navy = colors.HexColor(f"#{NAVY}")
    base = ParagraphStyle(
        "CLBody", parent=styles["Normal"], fontName=fonts["regular"],
        fontSize=10.5, leading=14, spaceAfter=9, textColor=colors.HexColor("#222222"),
    )
    name_style = ParagraphStyle(
        "CLName", parent=base, alignment=TA_CENTER, fontName=fonts["bold"],
        fontSize=18, leading=20, textColor=navy, spaceAfter=1,
    )
    contact_style = ParagraphStyle(
        "CLContact", parent=base, alignment=TA_CENTER, fontSize=9.4, leading=11,
        textColor=gray, spaceAfter=2,
    )

    sep = "&nbsp;&nbsp;|&nbsp;&nbsp;"
    parts = [escape(source.email), escape(source.phone)]
    if source.github:
        parts.append(f'<link href="{escape(source.github, quote=True)}" color="#{BLUE}">GitHub</link>')
    if source.linkedin:
        parts.append(f'<link href="{escape(source.linkedin, quote=True)}" color="#{BLUE}">LinkedIn</link>')
    contacts = sep.join(part for part in parts if part)
    story: list = [
        Paragraph(escape(source.name), name_style),
        Paragraph(contacts, contact_style),
        HRFlowable(width="100%", thickness=0.6, spaceBefore=1, spaceAfter=10,
                   color=colors.HexColor("#9AA5B1")),
        Paragraph(date.today().strftime("%B %-d, %Y") if os.name != "nt"
                  else date.today().strftime("%B %#d, %Y"), base),
        Spacer(1, 6),
    ]
    for para in re.split(r"\n\s*\n", body_text.strip()):
        if para.strip():
            story.append(Paragraph(escape(para.strip()).replace("\n", "<br/>"), base))
    document.build(story)
    return path


def _pdf_fonts() -> dict[str, str]:
    """Register Calibri from the system font directory if available.

    Returns a family map for regular/bold/italic/bolditalic. Falls back to the
    built-in Helvetica family when Calibri is not installed, so PDF generation
    never depends on a specific machine.
    """

    global _pdf_fonts_cache
    if _pdf_fonts_cache is not None:
        return _pdf_fonts_cache

    fonts = dict(_HELVETICA_FONTS)
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        fonts_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
        files = {
            "Calibri": "calibri.ttf", "Calibri-Bold": "calibrib.ttf",
            "Calibri-Italic": "calibrii.ttf", "Calibri-BoldItalic": "calibriz.ttf",
        }
        paths = {name: fonts_dir / filename for name, filename in files.items()}
        if all(path.exists() for path in paths.values()):
            for font_name, path in paths.items():
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            pdfmetrics.registerFontFamily(
                "Calibri", normal="Calibri", bold="Calibri-Bold",
                italic="Calibri-Italic", boldItalic="Calibri-BoldItalic",
            )
            fonts = {
                "regular": "Calibri", "bold": "Calibri-Bold",
                "italic": "Calibri-Italic", "bolditalic": "Calibri-BoldItalic",
            }
    except Exception:
        pass

    _pdf_fonts_cache = fonts
    return fonts


class ResumeDocumentGenerator:
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir

    def generate(self, resume: TailoredResume, job_id: int) -> tuple[Path, Path]:
        target = self.output_dir / f"{job_id}-{_slug(resume.target_company)}-{_slug(resume.target_title)}"
        target.mkdir(parents=True, exist_ok=True)
        stem = resume_file_stem(resume.source.name)
        docx_path = target / f"{stem}.docx"
        pdf_path = target / f"{stem}.pdf"
        self._build_docx(resume, docx_path)
        self._build_pdf(resume, pdf_path)
        return docx_path, pdf_path

    def _build_docx(self, resume: TailoredResume, path: Path) -> None:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.58)
        section.right_margin = Inches(0.58)
        section.header_distance = Inches(0.25)
        section.footer_distance = Inches(0.25)

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal.font.size = Pt(8.9)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0.6)
        normal.paragraph_format.line_spacing = 1.0

        heading = styles["Heading 1"]
        heading.font.name = "Calibri"
        heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        heading.font.size = Pt(10.2)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor.from_string(BLUE)
        heading.paragraph_format.space_before = Pt(3.2)
        heading.paragraph_format.space_after = Pt(1.4)
        heading.paragraph_format.keep_with_next = True

        bullet = styles["List Bullet"]
        bullet.font.name = "Calibri"
        bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        bullet.font.size = Pt(8.6)
        bullet.paragraph_format.left_indent = Inches(0.18)
        bullet.paragraph_format.first_line_indent = Inches(-0.13)
        bullet.paragraph_format.space_before = Pt(0)
        bullet.paragraph_format.space_after = Pt(0.35)
        bullet.paragraph_format.line_spacing = 1.0

        if "Entry" not in styles:
            entry = styles.add_style("Entry", WD_STYLE_TYPE.PARAGRAPH)
        else:
            entry = styles["Entry"]
        entry.base_style = normal
        entry.paragraph_format.space_before = Pt(0.6)
        entry.paragraph_format.space_after = Pt(0)
        entry.paragraph_format.keep_with_next = True

        def set_run(run, *, size=None, bold=None, italic=None, color=None):
            run.font.name = "Calibri"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
            if size is not None:
                run.font.size = Pt(size)
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if color is not None:
                run.font.color.rgb = RGBColor.from_string(color)

        def add_hyperlink(paragraph, label: str, url: str):
            part = paragraph.part
            relationship_id = part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), relationship_id)
            run = OxmlElement("w:r")
            properties = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), BLUE)
            properties.append(color)
            run.append(properties)
            text = OxmlElement("w:t")
            text.text = label
            run.append(text)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)

        def add_heading(text: str):
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.add_run(text.upper())
            return paragraph

        def set_right_tab(paragraph):
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(7.25), alignment=WD_TAB_ALIGNMENT.RIGHT
            )

        # Adapted proposal_centerpiece title block.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(resume.source.name), size=19.5, bold=True, color=NAVY)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1.2)
        set_run(p.add_run(resume.headline), size=9.4, italic=True, color=GRAY)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2.4)
        add_hyperlink(p, "GitHub", resume.source.github)
        set_run(p.add_run("  |  "), size=8.4, color=GRAY)
        add_hyperlink(p, "LinkedIn", resume.source.linkedin)
        set_run(p.add_run(f"  |  {resume.source.email}  |  {resume.source.phone}"), size=8.4)

        add_heading("Summary")
        p = doc.add_paragraph(resume.summary)
        p.paragraph_format.space_after = Pt(0.8)

        add_heading("Education")
        p = doc.add_paragraph(style="Entry")
        set_right_tab(p)
        set_run(p.add_run(resume.source.education.school), bold=True)
        p.add_run("\t")
        set_run(p.add_run(resume.source.education.location), italic=True)
        p = doc.add_paragraph(style="Entry")
        set_right_tab(p)
        set_run(p.add_run(resume.source.education.degree), italic=True)
        p.add_run("\t")
        set_run(p.add_run(resume.source.education.dates), italic=True)
        for detail in resume.source.education.details:
            doc.add_paragraph(detail, style="List Bullet")

        add_heading("Technical Skills")
        for group in resume.skill_groups:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0.25)
            set_run(p.add_run(f"{group.label}: "), bold=True)
            p.add_run(", ".join(group.skills))

        add_heading("Projects")
        for project in resume.projects:
            self._add_docx_project(doc, project, set_run, add_hyperlink, set_right_tab)

        add_heading("Experience")
        for item in resume.experience:
            p = doc.add_paragraph(style="Entry")
            set_right_tab(p)
            set_run(p.add_run(item.title), bold=True)
            p.add_run("\t")
            set_run(p.add_run(item.dates), italic=True)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(f"{item.organization} | {item.location}"), italic=True, color=GRAY)
            for bullet_text in item.bullets:
                doc.add_paragraph(bullet_text, style="List Bullet")

        doc.core_properties.author = resume.source.name
        doc.core_properties.title = f"Resume - {resume.target_title} - {resume.target_company}"
        doc.save(path)

    @staticmethod
    def _add_docx_project(doc, project: Project, set_run, add_hyperlink, set_right_tab) -> None:
        from docx.shared import Pt

        p = doc.add_paragraph(style="Entry")
        set_right_tab(p)
        set_run(p.add_run(project.title), bold=True)
        for label, url in project.links.items():
            set_run(p.add_run("  |  "), size=8.2, color=GRAY)
            add_hyperlink(p, label, url)
        p.add_run("\t")
        set_run(p.add_run(project.dates), italic=True)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(", ".join(project.skills)), italic=True, color=GRAY)
        for bullet_text in project.bullets:
            doc.add_paragraph(bullet_text, style="List Bullet")

    def _build_pdf(self, resume: TailoredResume, path: Path) -> None:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Table, TableStyle,
        )

        margin = 0.6 * inch
        page_width = letter[0]
        content_width = page_width - 2 * margin

        document = SimpleDocTemplate(
            str(path), pagesize=letter,
            leftMargin=margin, rightMargin=margin,
            topMargin=0.5 * inch, bottomMargin=0.5 * inch,
            title=f"Resume - {resume.target_title} - {resume.target_company}",
            author=resume.source.name,
        )
        styles = getSampleStyleSheet()
        fonts = _pdf_fonts()
        ink = colors.HexColor("#222222")
        gray = colors.HexColor(f"#{GRAY}")
        navy = colors.HexColor(f"#{NAVY}")
        link_blue = colors.HexColor(f"#{BLUE}")

        body = ParagraphStyle(
            "ResumeBody", parent=styles["Normal"], fontName=fonts["regular"],
            fontSize=9.6, leading=11.7, spaceAfter=0, textColor=ink,
        )
        heading = ParagraphStyle(
            "ResumeHeading", parent=body, fontName=fonts["bold"], fontSize=10.4,
            leading=11.8, spaceBefore=6, spaceAfter=2, textColor=navy, keepWithNext=True,
        )
        entry_left = ParagraphStyle("EntryLeft", parent=body, spaceAfter=0)
        entry_right = ParagraphStyle(
            "EntryRight", parent=body, fontSize=9.1, alignment=TA_RIGHT, textColor=gray,
        )
        subtitle = ParagraphStyle(
            "Subtitle", parent=body, fontSize=9.1, leading=11, textColor=gray, spaceAfter=1.2,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=body, fontSize=9.1, leading=11.2, spaceAfter=1.3,
            leftIndent=13, bulletIndent=2,
        )
        summary_style = ParagraphStyle("Summary", parent=body, spaceBefore=1, spaceAfter=1)
        name_style = ParagraphStyle(
            "Name", parent=body, alignment=TA_CENTER, fontName=fonts["bold"],
            fontSize=21, leading=23, textColor=navy, spaceAfter=1,
        )
        headline_style = ParagraphStyle(
            "Headline", parent=body, alignment=TA_CENTER, fontSize=10.5, leading=12.5,
            textColor=link_blue, spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "Contact", parent=body, alignment=TA_CENTER, fontSize=9.2, leading=11,
            textColor=gray, spaceAfter=2,
        )

        _cell_pad = [
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]

        def section(title: str) -> list:
            return [
                Paragraph(title.upper(), heading),
                HRFlowable(
                    width="100%", thickness=0.6, spaceBefore=0.5, spaceAfter=3,
                    color=colors.HexColor("#9AA5B1"),
                ),
            ]

        def header_row(left_html: str, right_text: str) -> Table:
            table = Table(
                [[Paragraph(left_html, entry_left), Paragraph(escape(right_text), entry_right)]],
                colWidths=[content_width * 0.74, content_width * 0.26],
            )
            table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "BOTTOM"), *_cell_pad]))
            return table

        def bullet_paras(items: tuple[str, ...]) -> list:
            return [Paragraph(escape(item), bullet_style, bulletText="•") for item in items]

        edu = resume.source.education
        story: list = [Paragraph(escape(resume.source.name), name_style)]
        if resume.headline:
            story.append(Paragraph(escape(resume.headline), headline_style))
        sep = "&nbsp;&nbsp;|&nbsp;&nbsp;"
        story.append(Paragraph(
            f'<link href="{escape(resume.source.github, quote=True)}" color="#{BLUE}">GitHub</link>{sep}'
            f'<link href="{escape(resume.source.linkedin, quote=True)}" color="#{BLUE}">LinkedIn</link>{sep}'
            f'{escape(resume.source.email)}{sep}{escape(resume.source.phone)}',
            contact_style,
        ))

        story += section("Summary")
        story.append(Paragraph(escape(resume.summary), summary_style))

        story += section("Education")
        story.append(header_row(f"<b>{escape(edu.school)}</b>", edu.location))
        story.append(header_row(f"<i>{escape(edu.degree)}</i>", edu.dates))
        story += bullet_paras(edu.details)

        story += section("Technical Skills")
        skill_rows = [
            [Paragraph(f"<b>{escape(group.label)}</b>", body),
             Paragraph(escape(", ".join(group.skills)), body)]
            for group in resume.skill_groups
        ]
        skills_table = Table(skill_rows, colWidths=[1.15 * inch, content_width - 1.15 * inch])
        skills_table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 1.2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1.2),
        ]))
        story.append(skills_table)

        story += section("Projects")
        for project in resume.projects:
            links = "".join(
                f'&nbsp;&nbsp;<link href="{escape(url, quote=True)}" color="#{BLUE}">{escape(label)}</link>'
                for label, url in project.links.items()
            )
            story.append(KeepTogether([
                header_row(f"<b>{escape(project.title)}</b>{links}", project.dates),
                Paragraph(f'<i>{escape(", ".join(project.skills))}</i>', subtitle),
                *bullet_paras(project.bullets),
            ]))

        story += section("Experience")
        for item in resume.experience:
            story.append(KeepTogether([
                header_row(f"<b>{escape(item.title)}</b>", item.dates),
                Paragraph(f"<i>{escape(item.organization)}{sep}{escape(item.location)}</i>", subtitle),
                *bullet_paras(item.bullets),
            ]))

        document.build(story)
